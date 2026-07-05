from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
font.size = Pt(10.5)
font.color.rgb = RGBColor(30, 30, 30)

style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

title = doc.add_heading('股票技术指标计算方法与作用说明', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].font.bold = True
title.runs[0].font.name = '宋体'
title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Created by Eric Zhao')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

indicators = [
    {
        'name': '一、均线系统（MA）',
        'full_name': 'Moving Average',
        'methods': [
            {'name': 'MA5（5日均线）', 'formula': 'MA5 = (C₁ + C₂ + C₃ + C₄ + C₅) / 5', 'desc': '最近5个交易日收盘价的算术平均值'},
            {'name': 'MA10（10日均线）', 'formula': 'MA10 = (C₁ + C₂ + ... + C₁₀) / 10', 'desc': '最近10个交易日收盘价的算术平均值'},
            {'name': 'MA20（20日均线）', 'formula': 'MA20 = (C₁ + C₂ + ... + C₂₀) / 20', 'desc': '最近20个交易日收盘价的算术平均值'},
            {'name': 'MA60（60日均线）', 'formula': 'MA60 = (C₁ + C₂ + ... + C₆₀) / 60', 'desc': '最近60个交易日收盘价的算术平均值'}
        ],
        'function': '均线系统用于判断趋势方向。短期均线反映近期市场情绪，中长期均线代表趋势方向。当短期均线上穿中长期均线形成"金叉"时通常是买入信号；反之"死叉"则预示调整。均线多头排列（MA5 > MA10 > MA20 > MA60）表示上涨趋势，空头排列则表示下跌趋势。',
        'strategies': ['趋势跟随策略', '支撑与阻力判断', '金叉死叉信号']
    },
    {
        'name': '二、MACD指标',
        'full_name': 'Moving Average Convergence Divergence',
        'methods': [
            {'name': 'EMA12（12日指数移动平均）', 'formula': 'EMA₁₂ = α × C + (1-α) × EMA₁₂(前)', 'desc': 'α = 2/(12+1) = 2/13 ≈ 0.1538'},
            {'name': 'EMA26（26日指数移动平均）', 'formula': 'EMA₂₆ = α × C + (1-α) × EMA₂₆(前)', 'desc': 'α = 2/(26+1) = 2/27 ≈ 0.0741'},
            {'name': 'DIF（差离值）', 'formula': 'DIF = EMA₁₂ - EMA₂₆', 'desc': '快线，反映短期与中期均线的差值'},
            {'name': 'DEA（异同平均数）', 'formula': 'DEA = EMA₉(DIF)', 'desc': '慢线，对DIF进行9日指数平滑'},
            {'name': 'MACD柱', 'formula': 'MACD柱 = DIF - DEA', 'desc': '本站采用不乘2的计算方式，直接等于DIF减DEA'}
        ],
        'function': 'MACD用于判断趋势强弱和多空转换。DIF上穿DEA形成金叉为买入信号，下穿形成死叉为卖出信号。MACD柱由负转正表示多头力量增强，由正转负表示空头力量增强。零轴上方表示多头趋势，零轴下方表示空头趋势。',
        'strategies': ['趋势跟随策略', '零轴过滤策略', '动量加速/减速策略', '多指标确认']
    },
    {
        'name': '三、RSI相对强弱指标',
        'full_name': 'Relative Strength Index',
        'methods': [
            {'name': '平均上涨幅度', 'formula': 'AvgUp = Wilder平滑(当日涨幅)', 'desc': '采用Wilder平滑方法，α = 1/14'},
            {'name': '平均下跌幅度', 'formula': 'AvgDown = Wilder平滑(当日跌幅)', 'desc': '采用Wilder平滑方法，α = 1/14'},
            {'name': 'RSI(14)', 'formula': 'RSI = 100 × AvgUp / (AvgUp + AvgDown)', 'desc': '将相对强弱压缩到0-100区间'}
        ],
        'function': 'RSI衡量市场上涨与下跌的相对力量。70为超买观察线，30为超卖观察线。RSI从低位向上穿越30关注反弹，从高位回落至70关注回调。RSI既可用于反转逻辑也可用于趋势强度逻辑，指标本身不是策略。',
        'strategies': ['均值回归策略', '趋势过滤策略', '背离研究']
    },
    {
        'name': '四、KDJ随机指标',
        'full_name': 'Stochastic Oscillator',
        'methods': [
            {'name': 'RSV（未成熟随机值）', 'formula': 'RSV = (C - L₉) / (H₉ - L₉) × 100', 'desc': 'C=收盘价，H₉=9日最高价，L₉=9日最低价'},
            {'name': 'K值', 'formula': 'K = EMA₃(RSV)', 'desc': '对RSV进行3日指数平滑，α=1/3（等价于SMA(RSV,3,1)）'},
            {'name': 'D值', 'formula': 'D = EMA₃(K)', 'desc': '对K值进行3日指数平滑，α=1/3（等价于SMA(K,3,1)）'},
            {'name': 'J值', 'formula': 'J = 3K - 2D', 'desc': '方向敏感线，反映K与D的差值'}
        ],
        'function': 'KDJ通过计算价格在高低区间的相对位置来判断超买超卖。80以上为超买区，20以下为超卖区。K线从下向上突破D线形成金叉且位于低位时是短期买入信号；K线从上向下跌破D线形成死叉且位于高位时是短期卖出信号。KDJ在震荡市中信号较为灵敏。',
        'strategies': ['超买超卖判断', '金叉死叉信号', 'J值预警']
    },
    {
        'name': '五、布林带（BOLL）',
        'full_name': 'Bollinger Bands',
        'methods': [
            {'name': '中轨（MA20）', 'formula': '中轨 = MA₂₀', 'desc': '20日简单移动平均线'},
            {'name': '上轨', 'formula': '上轨 = MA₂₀ + 2 × σ', 'desc': '中轨加上2倍标准差（总体标准差，ddof=0）'},
            {'name': '下轨', 'formula': '下轨 = MA₂₀ - 2 × σ', 'desc': '中轨减去2倍标准差（总体标准差，ddof=0）'},
            {'name': '带宽', 'formula': '带宽 = (上轨 - 下轨) / 中轨', 'desc': '反映波动率大小'}
        ],
        'function': '布林带反映价格的波动率和运行区间。布林带开口扩大表明波动率增加，通常伴随趋势加速；布林带收窄表明波动率降低，通常是变盘前兆。价格触及上轨时短期有回调压力，触及下轨时短期有反弹可能。同样触及上轨，均值回归策略可能考虑反向，突破策略可能考虑顺势。',
        'strategies': ['均值回归策略', '突破策略', '波动率收缩策略(Squeeze)', '动态区间/止损参考']
    },
    {
        'name': '六、ATR平均真实波幅',
        'full_name': 'Average True Range',
        'methods': [
            {'name': 'TR（真实波幅）', 'formula': 'TR = max(H-L, |H-Cₚ|, |L-Cₚ|)', 'desc': 'H=最高价，L=最低价，Cₚ=前收盘价'},
            {'name': 'ATR(14)', 'formula': 'ATR = Wilder平滑(TR, 14)', 'desc': '采用Wilder平滑方法，α = 1/14'}
        ],
        'function': 'ATR衡量价格波动幅度，不判断方向。ATR越大说明单日价格摆动越大，市场波动越剧烈；ATR越小说明波动越平静。ATR主要用于设置动态止损和仓位管理，ATR大时减少持仓数量，ATR小时适度增加仓位。',
        'strategies': ['波动率仓位管理', '动态止损策略', '突破过滤策略', '波动率目标/标的比较']
    }
]

for indicator in indicators:
    h1 = doc.add_heading(indicator['name'], level=1)
    h1.runs[0].font.name = '宋体'
    h1.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h1.runs[0].font.size = Pt(14)
    h1.runs[0].font.bold = True

    p = doc.add_paragraph(f'英文全称：{indicator["full_name"]}')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h2 = doc.add_heading('计算方法', level=2)
    h2.runs[0].font.name = '宋体'
    h2.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h2.runs[0].font.size = Pt(12)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标名称'
    hdr_cells[1].text = '计算公式'
    hdr_cells[2].text = '说明'
    for cell in hdr_cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for method in indicator['methods']:
        row_cells = table.add_row().cells
        row_cells[0].text = method['name']
        row_cells[1].text = method['formula']
        row_cells[2].text = method['desc']
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h2 = doc.add_heading('指标作用', level=2)
    h2.runs[0].font.name = '宋体'
    h2.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h2.runs[0].font.size = Pt(12)
    p = doc.add_paragraph(indicator['function'])
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h2 = doc.add_heading('策略角色', level=2)
    h2.runs[0].font.name = '宋体'
    h2.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h2.runs[0].font.size = Pt(12)
    p = doc.add_paragraph()
    for i, strategy in enumerate(indicator['strategies']):
        if i > 0:
            p.add_run('、')
        run = p.add_run(strategy)
        run.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

doc.save(r'D:\北大光华AI交易课程\赵展鹏-TASK2.docx')
print('文件已保存到：D:\\北大光华AI交易课程\\赵展鹏-TASK2.docx')