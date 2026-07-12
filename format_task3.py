from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document(r'd:\北大光华AI交易课程\赵展鹏-TASK3.docx')

for para in doc.paragraphs:
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)

doc.save(r'd:\北大光华AI交易课程\赵展鹏-TASK3.docx')
print('格式修改完成')
